import csv
import logging
import os
import sys
from abc import ABC, abstractmethod
from typing import List, Dict, Any

# Global imports (check for availability in __init__ and handle ImportError)
try:
    import docx
except ImportError:
    docx = None  # Mark as unavailable

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from PIL import Image
except ImportError:
    Image = None

try:
    import pytesseract
except ImportError:
    pytesseract = None

try:
    import fitz  # PyMuPDF is imported as fitz
except ImportError:
    fitz = None

try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None

from ..core.config import settings  # Example if config is needed
from ..core.exceptions import DocumentIngestionError  # Import custom exception

logger = logging.getLogger(__name__)


class BaseDocumentIngestor(ABC):
    """
    Abstract base class for all document ingestors.
    Defines the core interface for loading and extracting data from documents.
    """

    def __init__(self, file_path: str, config: Dict[str, Any] = None):
        """
        Initializes the ingestor with the file path and configuration options.

        Args:
            file_path (str): The path to the document file.
            config (Dict[str, Any]): A dictionary of file-type-specific options.
        """
        self.file_path = file_path
        self.config = config or {}
        if not os.path.exists(self.file_path):
            raise DocumentIngestionError(f"File not found: {self.file_path}")

    @abstractmethod
    def load_document(self) -> str:
        """
        Abstract method to load the document content. Subclasses must implement this.

        Returns:
            str: The raw text content of the document.
        """
        pass

    @abstractmethod
    def extract_metadata(self) -> Dict[str, Any]:
        """
        Abstract method to extract metadata from the document. Subclasses must implement this.

        Returns:
            Dict[str, Any]: A dictionary of metadata extracted from the document.
        """
        pass

    def ingest_document(self) -> Dict[str, Any]:
        """
        Combines load_document and extract_metadata into a single operation.

        Returns:
            Dict[str, Any]: A dictionary with doc_id, text, and metadata.
        """
        try:
            text = self.load_document()
            metadata = self.extract_metadata()
            doc_id = metadata.get("doc_id") or os.path.basename(self.file_path)

            return {"doc_id": str(doc_id), "text": text, "metadata": metadata}
        except Exception as e:
            msg = f"Error ingesting document {self.file_path}"
            logger.error(f"{msg}: {e}", exc_info=True)
            raise DocumentIngestionError(message=msg, details=str(e))


class CSVIngestor(BaseDocumentIngestor):
    """
    Handles CSV files for general-purpose use.
    It treats the entire CSV file as a single document, concatenating all cell
    content into one text block. This is ideal for RAG applications where
    the CSV's structure is unknown beforehand.
    """

    def __init__(self, file_path: str, config: Dict[str, Any] = None):
        """
        Initializes the general-purpose CSV ingestor.

        Args:
            file_path (str): The path to the CSV file.
            config (Dict[str, Any]): A dictionary of configuration options, including:
                encoding (str): The encoding of the CSV file (default: utf-8).
        """
        super().__init__(file_path, config)
        self.encoding = self.config.get("encoding", "utf-8")
        # Cache results to avoid reading the file twice
        self._full_text_content: str | None = None
        self._extracted_metadata: Dict[str, Any] | None = None

    def _read_and_parse_csv(self):
        """
        A private helper method to read the CSV file once, extracting both
        text and metadata in a single pass for efficiency.
        """
        # If we have already parsed the file, do nothing.
        if self._full_text_content is not None:
            return

        logger.info(f"Starting general ingestion for CSV: {self.file_path}")
        all_text_parts = []
        row_count = 0
        headers = []

        try:
            # Set a large field size limit for potentially large cells
            max_int = sys.maxsize
            while True:
                try:
                    csv.field_size_limit(max_int)
                    break
                except OverflowError:
                    max_int = int(max_int / 10)

            with open(self.file_path, mode="r", encoding=self.encoding, newline='') as csvfile:
                # Use the basic csv.reader as we make no assumptions about headers
                reader = csv.reader(csvfile)
                
                # Try to read the header row
                try:
                    headers = next(reader)
                    # Include headers in the text content for the LLM to have context
                    all_text_parts.append(", ".join(headers))
                except StopIteration:
                    # This means the file is empty
                    logger.warning(f"CSV file '{self.file_path}' is empty.")
                    self._full_text_content = ""
                    self._extracted_metadata = {
                        "filename": os.path.basename(self.file_path),
                        "doc_id": os.path.splitext(os.path.basename(self.file_path))[0],
                        "column_headers": [],
                        "row_count": 0,
                        "source_type": "csv"
                    }
                    return

                # Read the rest of the rows
                for row in reader:
                    # Join all non-empty cells in the row with a comma
                    row_text = ", ".join(cell.strip() for cell in row if cell and cell.strip())
                    if row_text:
                        all_text_parts.append(row_text)
                        row_count += 1
            
            # Combine all parts into a single text block
            self._full_text_content = "\n".join(all_text_parts)
            
            # Store the extracted metadata
            self._extracted_metadata = {
                "filename": os.path.basename(self.file_path),
                "doc_id": os.path.splitext(os.path.basename(self.file_path))[0],
                "column_headers": headers,
                "row_count": row_count,
            }
            logger.info(f"Successfully ingested CSV '{self.file_path}' with {len(headers)} columns and {row_count} data rows.")

        except FileNotFoundError as e:
            msg = f"CSV file not found at path: {self.file_path}"
            logger.error(msg)
            raise DocumentIngestionError(message=msg, details=str(e))
        except Exception as e:
            msg = f"An unexpected error occurred during general CSV ingestion from {self.file_path}"
            logger.error(f"{msg}: {e}", exc_info=True)
            raise DocumentIngestionError(message=msg, details=str(e))

    def load_document(self) -> str:
        """
        Loads the entire content of the CSV file as a single string.
        The first line of the output string will be the comma-separated headers,
        and subsequent lines will be the comma-separated text from each row.
        """
        self._read_and_parse_csv()
        return self._full_text_content

    def extract_metadata(self) -> Dict[str, Any]:
        """
        Extracts metadata about the CSV file, such as filename, column headers,
        and the number of rows.
        """
        self._read_and_parse_csv()
        return self._extracted_metadata

    # NOTE: The default ingest_document() method from the BaseDocumentIngestor
    # does not need to be overridden. It will work correctly with this new design.


class TXTIngestor(BaseDocumentIngestor):
    """
    Handles plain text files.
    Reads the entire text file and extracts basic metadata.
    """

    def __init__(self, file_path: str, config: Dict[str, Any] = None):
        """
        Initializes the TXT ingestor.

        Args:
            file_path (str): The path to the text file.
            config (Dict[str, Any]): A dictionary of configuration options, including:
                encoding (str): The encoding of the text file (default: utf-8).
        """
        super().__init__(file_path, config)
        self.encoding = self.config.get("encoding", "utf-8")

    def load_document(self) -> str:
        """Reads the entire text file."""
        try:
            with open(self.file_path, "r", encoding=self.encoding) as f:
                return f.read()
        except Exception as e:
            msg = f"Error reading TXT file: {self.file_path}"
            logger.error(f"{msg}: {e}", exc_info=True)
            raise DocumentIngestionError(message=msg, details=str(e))

    def extract_metadata(self) -> Dict[str, Any]:
        """Extracts basic metadata from the text file."""
        try:
            metadata = {
                "filename": os.path.basename(self.file_path),
                "last_modified": os.path.getmtime(self.file_path),
                "doc_id": os.path.splitext(os.path.basename(self.file_path))[0],
                "source_type": "txt"
            }
            return metadata
        except Exception as e:
            msg = f"Error extracting metadata from TXT file: {self.file_path}"
            logger.error(f"{msg}: {e}", exc_info=True)
            raise DocumentIngestionError(message=msg, details=str(e))


class PDFIngestor(BaseDocumentIngestor):
    """
    Handles PDF files, with robust support for both text-based and scanned (image-based) PDFs.
    """
    def __init__(self, file_path: str, config: Dict[str, Any] = None):
        super().__init__(file_path, config)
        self.ocr_language = self.config.get("ocr_language", "eng")
        if fitz is None:
            raise DocumentIngestionError("PyMuPDF library not found. Please install it.")

    def load_document(self) -> str:
        """
        Extracts text from a PDF. It first attempts a direct text extraction.
        If the extracted text is minimal or empty (indicating a scanned PDF),
        it automatically falls back to performing OCR on each page.
        """
        all_text = []
        try:
            # Attempt Direct Text Extraction
            logger.info(f"Attempting direct text extraction for PDF: {self.file_path}")
            doc = fitz.open(self.file_path)
            for page_num, page in enumerate(doc):
                page_text = page.get_text().strip()
                if page_text:
                    all_text.append(page_text)
            
            full_direct_text = "\n".join(all_text)

            # Check if Direct Extraction was Sufficient
            # If we got a reasonable amount of text, we can assume it's not a scanned document.
            if full_direct_text and (len(full_direct_text) > 10 * doc.page_count):
                logger.info("Direct text extraction successful. Skipping OCR.")
                return full_direct_text

            # Fallback to OCR if Needed
            logger.warning(f"Direct text extraction yielded minimal text for '{self.file_path}'. Falling back to OCR.")
            
            # Check for OCR dependencies
            if convert_from_path is None or pytesseract is None or Image is None:
                raise DocumentIngestionError("OCR dependencies (pdf2image, pytesseract, Pillow) are not installed, but the PDF appears to be scanned.")

            # Convert PDF pages to images
            try:
                images = convert_from_path(self.file_path)
            except Exception as e: # Catch poppler errors specifically if possible
                 if "Poppler" in str(e):
                      logger.error("Poppler utility not found. Please install it and ensure it's in your system's PATH to process scanned PDFs.")
                      raise DocumentIngestionError("Poppler not found. OCR on PDF is not possible without it.")
                 raise e


            ocr_text_content = []
            for i, image in enumerate(images):
                logger.debug(f"Performing OCR on page {i+1} of {len(images)}...")
                try:
                    text_from_page = pytesseract.image_to_string(image, lang=self.ocr_language)
                    if text_from_page:
                        ocr_text_content.append(text_from_page)
                except pytesseract.TesseractNotFoundError:
                    logger.error("Tesseract OCR engine not found. Ensure it's installed and in your PATH.")
                    raise DocumentIngestionError("Tesseract OCR engine not found.")

            logger.info(f"OCR processing completed for '{self.file_path}'.")
            return "\n".join(ocr_text_content)

        except Exception as e:
            msg = f"An unexpected error occurred during PDF processing for '{self.file_path}'"
            logger.error(f"{msg}: {e}", exc_info=True)
            raise DocumentIngestionError(message=msg, details=str(e))

    def extract_metadata(self) -> Dict[str, Any]:
        """Extracts metadata from the PDF file using PyMuPDF."""
        try:
            doc = fitz.open(self.file_path)
            
            # PyMuPDF metadata is already a standard Python dictionary
            clean_metadata = doc.metadata or {}

            # Add our custom metadata
            clean_metadata["filename"] = os.path.basename(self.file_path)
            clean_metadata["doc_id"] = os.path.splitext(os.path.basename(self.file_path))[0]
            clean_metadata["page_count"] = len(doc)
            clean_metadata["source_type"] = "pdf"
            
            doc.close()
            return clean_metadata
        except Exception as e:
            msg = f"Error extracting metadata from PDF file: {self.file_path}"
            logger.error(f"{msg}: {e}", exc_info=True)
            raise DocumentIngestionError(message=msg, details=str(e))

class DOCXIngestor(BaseDocumentIngestor):
    """
    Handles DOCX files.
    Uses python-docx to extract text and metadata.
    """

    def __init__(self, file_path: str, config: Dict[str, Any] = None):
        """
        Initializes the DOCX ingestor.

        Args:
            file_path (str): The path to the DOCX file.
            config (Dict[str, Any]): A dictionary of configuration options (currently none).
        """
        super().__init__(file_path, config)

    def load_document(self) -> str:
        """Extracts text from the DOCX file using python-docx."""
        if docx is None:
            logger.error("python-docx library is not installed. Cannot process DOCX files.")
            raise DocumentIngestionError(message="python-docx library not found. Install it to process DOCX files.")
        try:
            doc = docx.Document(self.file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            msg = f"Error reading DOCX file: {self.file_path}"
            logger.error(f"{msg}: {e}", exc_info=True)
            raise DocumentIngestionError(message=msg, details=str(e))

    def extract_metadata(self) -> Dict[str, Any]:
        """Extracts metadata from the DOCX file."""
        if docx is None:
            logger.error("python-docx library is not installed. Cannot process DOCX files.")
            raise DocumentIngestionError(message="python-docx library not found. Install it to process DOCX files.")
        try:
            doc = docx.Document(self.file_path)
            properties = doc.core_properties
            metadata = {
                "filename": os.path.basename(self.file_path),
                "author": properties.author,
                "title": properties.title,
                "created": str(properties.created), # Convert to string for easier handling
                "modified": str(properties.modified), # Convert to string
                "doc_id": os.path.splitext(os.path.basename(self.file_path))[0],
                "source_type": "docx"
            }
            return metadata
        except Exception as e:
            msg = f"Error extracting metadata from DOCX file: {self.file_path}"
            logger.error(f"{msg}: {e}", exc_info=True)
            raise DocumentIngestionError(message=msg, details=str(e))


class ImageIngestor(BaseDocumentIngestor):
    pass


class ExcelIngestor(BaseDocumentIngestor):
    pass


class DocumentIngestorFactory:
    """
    Creates the appropriate BaseDocumentIngestor subclass based on the file type.
    """

    def create_ingestor(self, file_path: str, config: Dict[str, Any] = None) -> BaseDocumentIngestor:
        """
        Determines the file type and instantiates the correct ingestor class.

        Args:
            file_path (str): The path to the document file.
            config (Dict[str, Any]): A dictionary of file-type-specific options.

        Returns:
            BaseDocumentIngestor: An instance of the appropriate ingestor class.

        Raises:
            ValueError: If the file type is not supported.
        """
        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension == ".csv":
            return CSVIngestor(file_path, config)
        elif file_extension == ".txt":
            return TXTIngestor(file_path, config)
        elif file_extension == ".pdf":
            return PDFIngestor(file_path, config)
        elif file_extension == ".docx":
            return DOCXIngestor(file_path, config)
        elif file_extension in [".png", ".jpg", ".jpeg", ".tiff", ".tif"]:
            return ImageIngestor(file_path, config)
        elif file_extension in [".xlsx", ".xls"]:
            return ExcelIngestor(file_path, config)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

# For local testing 
if __name__ == "__main__":
    # Configure logging for the test run
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    logger.info("--- Starting Real File Ingestion Tests ---")
    
    # --- Setup Paths ---
    # This assumes the script is in 'backend/app/services/'
    # It navigates up to the 'backend' directory and then into 'data'
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        backend_dir = os.path.dirname(os.path.dirname(current_dir))
        data_dir = os.path.join(backend_dir, 'data')

        logger.info(f"Data directory set to: {data_dir}")

        if not os.path.isdir(data_dir):
            raise FileNotFoundError(f"The data directory was not found at the expected path: {data_dir}")

        # Define paths to the test files
        txt_test_path = os.path.join(data_dir, "sample_document.txt")
        csv_test_path = os.path.join(data_dir, "sample_data.csv")
        pdf_test_path = os.path.join(data_dir, "sample_report.pdf")
        docx_test_path = os.path.join(data_dir, "sample_document.docx")
        # Instantiate the factory
        factory = DocumentIngestorFactory()

        # --- Test DOCX Ingestion ---
        logger.info("\n--- Testing DOCX Ingestion ---")
        if os.path.exists(docx_test_path):
            try:
                docx_ingestor = factory.create_ingestor(docx_test_path)
                docx_data = docx_ingestor.ingest_document()
                logger.info(f"DOCX Ingestion successful. Doc ID: {docx_data.get('doc_id')}, Text Snippet: '{docx_data.get('text', '')[:70]}...'")
                # logger.debug(f"Full DOCX Data: {docx_data}")
            except Exception as e:
                logger.error(f"Failed to ingest DOCX file '{docx_test_path}': {e}", exc_info=True)

        # --- Test TXT Ingestion ---
        logger.info("\n--- Testing TXT Ingestion ---")
        if os.path.exists(txt_test_path):
            try:
                txt_ingestor = factory.create_ingestor(txt_test_path)
                txt_data = txt_ingestor.ingest_document()
                logger.info(f"TXT Ingestion successful. Doc ID: {txt_data.get('doc_id')}, Text Snippet: '{txt_data.get('text', '')[:70]}...'")
                # logger.debug(f"Full TXT Data: {txt_data}")
            except Exception as e:
                logger.error(f"Failed to ingest TXT file '{txt_test_path}': {e}", exc_info=True)
        else:
            logger.warning(f"Skipping TXT test: File not found at '{txt_test_path}'")

        # --- Test CSV Ingestion ---
        logger.info("\n--- Testing CSV Ingestion ---")
        if os.path.exists(csv_test_path):
            try:
                csv_ingestor = factory.create_ingestor(csv_test_path)
                csv_data = csv_ingestor.ingest_document()
                logger.info(f"CSV Ingestion successful (first row). Doc ID: {csv_data.get('doc_id')}, Text Snippet: '{csv_data.get('text', '')[:70]}...'")
                logger.info(f"Full CSV Data: {csv_data}")
            except Exception as e:
                logger.error(f"Failed to ingest CSV file '{csv_test_path}': {e}", exc_info=True)
        else:
            logger.warning(f"Skipping CSV test: File not found at '{csv_test_path}'")

        logger.info("\n--- Testing PDF Ingestion ---")
        if os.path.exists(pdf_test_path):
            try:
                pdf_ingestor = factory.create_ingestor(pdf_test_path)
                pdf_data = pdf_ingestor.ingest_document()
                logger.info(f"PDF Ingestion successful.\nDoc ID: {pdf_data.get('doc_id')}, Text Snippet: '{pdf_data.get('text', '').strip()[:700]}...'")
                #logger.info(f"\nFull PDF Data: {pdf_data}")
                logger.info(f"\nPDF Metadata: {pdf_data.get('metadata')}")
            except Exception as e:
                logger.error(f"Failed to ingest PDF file '{pdf_test_path}': {e}", exc_info=True)
        else:
            logger.warning(f"Skipping PDF test: File not found at '{pdf_test_path}'")

    except Exception as e:
        logger.error(f"A critical error occurred during test setup: {e}", exc_info=True)

    logger.info("\n--- End of Ingestion Tests ---")